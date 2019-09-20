"""Test Chapter 6."""
import os
import unittest.mock
from io import StringIO
from platform import system

from docx import Document
from docx.shared import RGBColor

import src.ch06.p1_invisible_ink as invisible_ink
import src.ch06.c1_invisible_ink_mono as invisible_ink_mono


class TestInvisibleInk(unittest.TestCase):
    """Test Invisible Ink."""

    def test_get_text(self):
        """Test get_text."""
        testfile = os.path.normpath('tests/data/ch06/fake.docx')
        # Test that it doesn't skip blanks.
        paragraphs = invisible_ink.get_text(testfile, skip_blank=False)
        self.assertEqual(paragraphs.count(''), 4)
        # Test that it does skip blanks.
        paragraphs = invisible_ink.get_text(testfile)
        self.assertEqual(paragraphs.count(''), 0)
        # Test that it read contents.
        answerlist = \
            ['This is a test document.',
             'This is a paragraph with two runs. However, it’s not because it has two '
             'lines.',
             'There is intentionally a lot of blank spaces to check if the code can count '
             'them correctly.',
             'So, don’t send me e-mails saying that the formatting in my test files is '
             'incorrect.',
             'Word.']
        self.assertEqual(answerlist, paragraphs)

    def test_check_blanks(self):
        """Test check_blanks."""
        fakefile = os.path.normpath('tests/data/ch06/fake.docx')
        cipherfile = os.path.normpath('tests/data/ch06/cipher.docx')
        # Test that it doesn't need extra blanks.
        faketext = invisible_ink.get_text(fakefile, False)
        ciphertext = invisible_ink.get_text(cipherfile)
        blanks_needed = invisible_ink.check_blanks(faketext, ciphertext)
        self.assertEqual(blanks_needed, 0)
        # Test that it does need extra blanks.
        faketext = invisible_ink.get_text(fakefile)
        blanks_needed = invisible_ink.check_blanks(faketext, ciphertext)
        self.assertEqual(blanks_needed, 3)
        faketext.append('')
        blanks_needed = invisible_ink.check_blanks(faketext, ciphertext)
        self.assertEqual(blanks_needed, 2)
        faketext.append('')
        blanks_needed = invisible_ink.check_blanks(faketext, ciphertext)
        self.assertEqual(blanks_needed, 1)
        faketext.append('')
        blanks_needed = invisible_ink.check_blanks(faketext, ciphertext)
        self.assertEqual(blanks_needed, 0)

    def test_write_invisible(self):
        """Test write_invisible."""
        fakefile = os.path.normpath('tests/data/ch06/fake.docx')
        cipherfile = os.path.normpath('tests/data/ch06/cipher.docx')
        faketext = invisible_ink.get_text(fakefile, False)
        ciphertext = invisible_ink.get_text(cipherfile)
        current_dir = os.path.curdir
        # Test default template and filename.
        invisible_ink.write_invisible(faketext, ciphertext)
        output_file = os.path.join(current_dir, 'output.docx')
        self.assertTrue(os.path.exists(output_file))
        output_text = invisible_ink.get_text(output_file)
        all_text = [element for element in faketext if element != ''] + \
            ciphertext
        self.assertEqual(len(all_text), len(output_text))
        for line in output_text:
            self.assertIn(line, all_text)
        os.remove(output_file)
        # Test custom template and filename.
        template_file = os.path.normpath('tests/data/ch06/template.docx')
        output_file = os.path.join(current_dir, 'letter.docx')
        invisible_ink.write_invisible(faketext, ciphertext, template_file, 'letter.docx')
        self.assertTrue(os.path.exists(output_file))
        output_text = invisible_ink.get_text(output_file)
        all_text = [element for element in faketext if element != ''] + \
            ciphertext
        self.assertEqual(len(all_text), len(output_text))
        for line in output_text:
            self.assertIn(line, all_text)
        os.remove(output_file)
        # Test error.
        faketext = invisible_ink.get_text(fakefile)
        error = ('3 more blanks are needed in the '
                 'plaintext (fake) message.')
        with self.assertRaises(ValueError) as err:
            invisible_ink.write_invisible(faketext, ciphertext)
        self.assertEqual(error, str(err.exception))

    @unittest.mock.patch('src.ch06.p1_invisible_ink.Path.resolve')
    @unittest.mock.patch('sys.stdout', new_callable=StringIO)
    def test_main(self, mock_stdout, mock_abspath):
        """Test demo main function."""
        # FIXME: Why doesn't current_dir = os.path.abspath('./p1files') and
        #   fakefile = os.path.join(current_dir, 'fake.docx') used in
        #   src/ch06.p1_invisible_ink.main work with
        #   @unittest.mock.patch('src.ch06.p1_invisible_ink.os.path.abspath')
        #   and mock_abspath.return_value = os.path.normpath('src/ch06/p1files')
        #   used here?
        #   Using Python 3.6.8 and python-docx 0.8.10, fails with:
        #   ValueError: PackURI must begin with slash, got 'src/ch06/p1files'
        #   Had to use pathlib's Path and PurePath in p1_invisible_ink.main

        # Mock output of abspath to avoid FileNotFoundError.
        mock_abspath.return_value = os.path.normpath('src/ch06/p1files')
        current_dir = os.getcwd()
        # Test using test files.
        fakefile = os.path.join(current_dir, 'tests/data/ch06/fake.docx')
        cipherfile = os.path.join(current_dir, 'tests/data/ch06/cipher.docx')
        output_file = os.path.join(current_dir, 'tests/data/ch06/output.docx')
        faketext = invisible_ink.get_text(fakefile, False)
        ciphertext = invisible_ink.get_text(cipherfile)
        invisible_ink.main(fakefile, cipherfile, output_file)
        self.assertTrue(os.path.exists(output_file))
        output_text = invisible_ink.get_text(output_file)
        all_text = ([element for element in faketext if element != ''] +
                    ciphertext)
        self.assertEqual(len(all_text), len(output_text))
        for line in output_text:
            self.assertIn(line, all_text)
        os.remove(output_file)
        # Test printed output.
        with open(os.path.normpath('tests/data/ch06/main/invisible_ink.txt'),
                  'r') as file:
            file_data = ''.join(file.readlines())
        self.assertEqual(mock_stdout.getvalue(), file_data)
        # Test using default files.
        invisible_ink.main()
        fakefile = os.path.join(current_dir, 'src/ch06/p1files/fake.docx')
        cipherfile = os.path.join(current_dir, 'src/ch06/p1files/real.docx')
        output_file = os.path.normpath('src/ch06/p1files/LetterToUSDA.docx')
        faketext = invisible_ink.get_text(fakefile, False)
        ciphertext = invisible_ink.get_text(cipherfile)
        self.assertTrue(os.path.exists(output_file))
        output_text = invisible_ink.get_text(output_file)
        all_text = ([element for element in faketext if element != ''] +
                    ciphertext)
        self.assertEqual(len(all_text), len(output_text))
        for line in output_text:
            self.assertIn(line, all_text)
        os.remove(output_file)


class TestInvisibleInkMono(unittest.TestCase):
    """Test Invisible Ink Mono."""

    def test_check_fit(self):
        """Test check_fit."""
        fakefile = os.path.normpath('tests/data/ch06/fake_mono.docx')
        cipherfile = os.path.normpath('tests/data/ch06/cipher_mono.docx')
        # Test that it doesn't need extra blanks.
        faketext = invisible_ink.get_text(fakefile, False)
        ciphertext = invisible_ink.get_text(cipherfile)
        blanks_needed = invisible_ink_mono.check_fit(faketext, ciphertext)
        self.assertEqual(blanks_needed, 0)
        # Test that it does need extra blanks.
        faketext = ['This is too short.']
        blanks_needed = invisible_ink_mono.check_fit(faketext, ciphertext)
        self.assertEqual(blanks_needed, 49)
        faketext.append('You would have to write a small novel to get it to '
                        'fit.')
        blanks_needed = invisible_ink_mono.check_fit(faketext, ciphertext)
        self.assertEqual(blanks_needed, 37)
        faketext.append('Filling in blanks is not as easy as it seems because '
                        'so few are in every sentence.')
        blanks_needed = invisible_ink_mono.check_fit(faketext, ciphertext)
        self.assertEqual(blanks_needed, 21)
        faketext.append('The use of small words helps, but it is not a good '
                        'way to go about being a super secret spy person.')
        blanks_needed = invisible_ink_mono.check_fit(faketext, ciphertext)
        self.assertEqual(blanks_needed, 0)

    def test_write_invisible(self):
        """Test write_invisible."""
        fakefile = os.path.normpath('tests/data/ch06/fake_mono.docx')
        cipherfile = os.path.normpath('tests/data/ch06/cipher_mono.docx')
        faketext = invisible_ink.get_text(fakefile, False)
        ciphertext = invisible_ink.get_text(cipherfile)
        current_dir = os.path.curdir
        # Test default template and filename.
        invisible_ink_mono.write_invisible(faketext, ciphertext)
        output_file = os.path.join(current_dir, 'output.docx')
        self.assertTrue(os.path.exists(output_file))
        output_text = invisible_ink.get_text(output_file)
        answer_text = [
            'ThisTishaitestsdocument withiaslot ofsfiller,'
            'uorpunnecessarypwordiness.',
            'Please,otrysnotetodwrite liketthisobecause itbiseas annoyingaas '
            'itsishunnecessary.',
            'Unless,oofrcourse,tyou,are '
            'writinguantypeeofncharactercthatrisywordypandtusesewordsd'
            'unnecessarily.',
            'In thatmrare,euncommonsinstance,sitaisgperfectlyepermissible.'
            'to be wordy.']
        self.assertListEqual(answer_text, output_text)
        # Check color
        paragraph_index, count = 0, 0
        cipher_len = sum(len(line) for line in ciphertext)
        doc = Document(output_file)
        while count < cipher_len:
            for line in faketext:
                paragraph = doc.paragraphs[paragraph_index]
                if line == '':
                    # Skip blanks in faketext and output_file.
                    paragraph_index += 1
                    continue
                letter_index = 0
                for word in line.split():
                    # Check color of each letter after word.
                    letter_index += len(word)
                    if letter_index >= len(line):
                        # Stop checking at the end of the line.
                        break
                    run = paragraph.runs[letter_index]
                    if all([len(run.text) == 1, run.text != ' ']):
                        self.assertEqual(run.font.color.rgb, RGBColor(255, 255, 255))
                    count += 1
                    letter_index += 1
                paragraph_index += 1
        os.remove(output_file)
        # Test custom template and filename.
        template_file = os.path.normpath('tests/data/ch06/template_mono.docx')
        output_file = os.path.join(current_dir, 'letter.docx')
        invisible_ink_mono.write_invisible(faketext, ciphertext, template_file, 'letter.docx')
        self.assertTrue(os.path.exists(output_file))
        output_text = invisible_ink.get_text(output_file)
        self.assertListEqual(answer_text, output_text)
        # Check color
        paragraph_index, count = 0, 0
        cipher_len = sum(len(line) for line in ciphertext)
        doc = Document(output_file)
        while count < cipher_len:
            for line in faketext:
                paragraph = doc.paragraphs[paragraph_index]
                if line == '':
                    # Skip blanks in faketext and output_file.
                    paragraph_index += 1
                    continue
                if paragraph.text == '':
                    # FIXME: template_file always has a blank paragraph.
                    paragraph_index += 1
                    paragraph = doc.paragraphs[paragraph_index]
                letter_index = 0
                for word in line.split():
                    # Check color of each letter after word.
                    letter_index += len(word)
                    if letter_index >= len(line):
                        # Stop checking at the end of the line.
                        break
                    run = paragraph.runs[letter_index]
                    if all([len(run.text) == 1, run.text != ' ']):
                        self.assertEqual(run.font.color.rgb, RGBColor(255, 255, 255))
                    count += 1
                    letter_index += 1
                paragraph_index += 1
        os.remove(output_file)
        # Test font name.
        invisible_ink_mono.write_invisible(faketext, ciphertext, None, 'letter.docx')
        doc = Document(output_file)
        if system().lower().startswith('windows'):
            for paragraph in doc.paragraphs:
                if paragraph.text == '':
                    continue
                self.assertEqual(paragraph.style.font.name, "Courier New")
        else:
            for paragraph in doc.paragraphs:
                if paragraph.text == '':
                    continue
                self.assertEqual(paragraph.style.font.name, "Liberation Mono")
        os.remove(output_file)
        # Test error.
        faketext = invisible_ink.get_text(fakefile)[2:]
        error = 'Need 25 more spaces in the plaintext (fake) message.'
        with self.assertRaises(ValueError) as err:
            invisible_ink_mono.write_invisible(faketext, ciphertext)
        self.assertEqual(error, str(err.exception))
        # TODO: Test multi-line ciphertext.


if __name__ == '__main__':
    unittest.main()
