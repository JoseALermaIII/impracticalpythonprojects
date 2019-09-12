"""Use stenography to hide messages in a word processor document.

Use :py:obj:`docx.Document` to hide encrypted messages in a word processor
document by embedding the encrypted message in a fake message's whitespace,
then changing the encrypted message's font color to white.

Note:
    Using LibreOffice version 6.0.7.3

Warning:
    There are many ways this method of stenography can fail. Please don't use
    for actual covert operations (covered in MIT License).

"""
from pathlib import Path, PurePath
import docx
from docx.shared import RGBColor, Pt


def get_text(file_path: str, skip_blank: bool = True) -> list:
    """Get text from a docx file.

    Loads paragraphs from the given docx file into a list. Optionally skips
    blank lines.

    Args:
        file_path (str): Absolute path to a .docx file to load.
        skip_blank (bool): Whether or not to skip blank lines. Defaults
            to :py:obj:`True`.
    Returns:
        Each paragraph in the docx file in a list of strings.

    Note:
        Does not copy formatting from docx file - only text.

    """
    paragraphs = []
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        if all([skip_blank, not paragraph.text]):
            continue
        paragraphs.append(paragraph.text)
    return paragraphs


def check_blanks(plaintext: list, ciphertext: list) -> int:
    """Check if the ciphertext can fit in plaintext.

    Compare the number of blank lines in **plaintext** to the number of lines
    in **ciphertext**. If they aren't a match, returns the number of extra
    blank lines needed.

    Args:
        plaintext (list): Paragraphs of a fake message in a list of strings
            (likely from :func:`get_text`).
        ciphertext (list): Paragraphs of an encrypted message in a list of
            strings (likely from :func:`get_text`).

    Returns:
        Integer representing the number of needed blank lines to fit
        **ciphertext** in **plaintext**. ``0`` would mean that **ciphertext**
        can fit in **plaintext**.

    """
    blanks_needed = len(ciphertext) - plaintext.count('')
    if blanks_needed <= 0:
        return 0
    return blanks_needed


def write_invisible(plaintext: list, ciphertext: list,
                    template_path: str = None,
                    filename: str = 'output.docx') -> None:
    """Embed ciphertext in plaintext's whitespace.

    Open a template file, **template_path**, with the needed fonts, styles,
    and margins. Write each paragraph in **plaintext** to the template file
    and add each paragraph in **ciphertext** to **plaintext**'s blank space.
    Save the new file as **filename**.

    Args:
        plaintext (list): Paragraphs of a fake message in a list of strings
            (likely from :func:`get_text`).
        ciphertext (list): Paragraphs of an encrypted message in a list of
            strings (likely from :func:`get_text`).
        template_path (str): Absolute path to .docx file with predefined
            fonts, styles, and margins. Defaults to :py:obj:`None`. If not
            provided, defaults will be created.
        filename (str): File name to use for output file. Defaults to
            ``output.docx``.

    Returns:
        :py:obj:`None`. **plaintext** is written to the file at
        **template_path** with **ciphertext** embedded in the blank space.

    Raises:
        ValueError: If the number of blank lines in **plaintext** aren't
            enough to embed **ciphertext** based on output of
            :func:`check_blanks`.

    Note:
        As of python-docx v0.8.10, creating custom styles isn't well
        supported. More info `here`_.

        As a result, if a template isn't provided, the default template is
        used.

    .. _here:
        https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html

    """
    blanks_needed = check_blanks(plaintext, ciphertext)
    if blanks_needed > 0:
        raise ValueError(f'{blanks_needed} more blanks are needed in the '
                         f'plaintext (fake) message.')
    if template_path is None:
        # Use default template.
        doc = docx.Document()
    else:
        doc = docx.Document(template_path)

    index = 0
    for line in plaintext:
        if all([line == '', index < len(ciphertext)]):
            paragraph = doc.add_paragraph(ciphertext[index])
            paragraph_index = len(doc.paragraphs) - 1
            # Set real message color to white.
            run = doc.paragraphs[paragraph_index].runs[0]
            font = run.font
            # Make red for testing.
            font.color.rgb = RGBColor(255, 255, 255)
            index += 1
        else:
            paragraph = doc.add_paragraph(line)
        # Set line spacing between paragraphs.
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    doc.save(filename)


def main(fakefile: str = None, cipherfile: str = None,
         savepath: str = None) -> None:
    """Demonstrate the invisible ink writer.

    Demonstrate :func:`write_invisible`, but for testing,
    it is a basic wrapper function for :func:`write_invisible`.
    Embed **cipherfile** in **fakefile**'s whitespace.

    Args:
        fakefile (str): Path to .docx file with fake message.
            Defaults to ``./p1files/fake.docx``.
        cipherfile (str): Path to .docx file with real message.
            Defaults to ``./p1files/real.docx``.
        savepath (str): Path to .docx file for output.
            Defaults to ``./p1files/LetterToUSDA.docx``.

    Returns:
         :py:obj:`None`. The contents of **cipherfile**'s text is embedded
         in **fakefile**'s whitespace and saved to **savepath**.

    """
    print('I can embed a hidden message in a .docx file\'s white space '
          'by making the font\ncolor white. It isn\'t as bulletproof '
          'as it sounds.\n')
    current_dir = Path('./p1files').resolve()
    if fakefile is None or cipherfile is None:
        fakefile = PurePath(current_dir).joinpath('fake.docx')
        cipherfile = PurePath(current_dir).joinpath('real.docx')
    if savepath is None:
        savepath = PurePath(current_dir).joinpath('LetterToUSDA.docx')
    faketext = get_text(fakefile, False)
    ciphertext = get_text(cipherfile)
    write_invisible(faketext, ciphertext, None, savepath)
    print('Done.\n')
    print('To read the real message, select the entire document and\n'
          'highlight it a dark gray.')


if __name__ == '__main__':
    main()
