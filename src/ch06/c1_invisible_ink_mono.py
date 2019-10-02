"""Use stenography to hide messages in a word processor document.

Use :py:obj:`docx.Document` to hide encrypted messages in a word processor
document by embedding the encrypted message in a fake message's whitespace,
then changing the encrypted message's font color to white.

Note:
    Using LibreOffice version 6.0.7.3

Warning:
    There are many ways this method of stenography can fail. Please don't use
    for actual covert operations (covered in MIT License).

"""
from pathlib import Path, PurePath
from platform import system

import docx
from docx.shared import RGBColor

from src.ch06.p1_invisible_ink import get_text


def check_fit(plaintext: list, ciphertext: list) -> int:
    """Check if ciphertext can fit in plaintext's whitespace.

    Sum number of blanks in **plaintext** and compare to number of characters
    in **ciphertext** to see if it can fit.

    Args:
        plaintext (list): Paragraphs of a fake message in a list of strings
            (likely from :func:`~src.ch06.p1_invisible_ink.get_text`).
        ciphertext (list): Paragraphs of an encrypted message in a list of
            strings (likely from :func:`~src.ch06.p1_invisible_ink.get_text`).

    Returns:
        Integer representing the number of needed blanks to fit
        **ciphertext** in **plaintext**. ``0`` would mean that **ciphertext**
        can fit in **plaintext**.

    Note:
        To separate words, the blanks in **ciphertext** count toward the
        needed length of **plaintext**. By contrast, blank lines in
        **plaintext** do not count.

    """
    blanks = sum(line.count(' ') for line in plaintext if line != '')
    letters = sum(len(line) for line in ciphertext if line != '')
    if blanks >= letters:
        return 0
    return letters - blanks


def write_invisible(plaintext: list, ciphertext: list,
                    template_path: str = None,
                    filename: str = 'output.docx') -> None:
    """Embed ciphertext in plaintext's letter whitespace.

    Open a template file, **template_path**, with the needed fonts, styles,
    and margins. Write each line in **plaintext** to the template file
    and add each line in **ciphertext** to **plaintext**'s space between
    letters by using a monospace font.
    Save the new file as **filename**.

    Args:
        plaintext (list): Lines of a fake message in a list of strings
            (likely from :func:`~src.ch06.p1_invisible_ink.get_text`).
        ciphertext (list): Lines of an encrypted message in a list of
            strings (likely from :func:`~src.ch06.p1_invisible_ink.get_text`).
        template_path (str): Absolute path to .docx file with predefined
            fonts, styles, and margins. Defaults to :py:obj:`None`. If not
            provided, defaults will be created.
        filename (str): File name to use for output file. Defaults to
            ``output.docx``.

    Returns:
        :py:obj:`None`. **plaintext** is written to the file at
        **template_path** with **ciphertext** embedded in the blank space.

    Raises:
        ValueError: If the number of spaces in **plaintext** aren't
            enough to embed **ciphertext** based on output of
            :func:`check_fit`.

    Note:
        As of python-docx v0.8.10, creating custom styles isn't well
        supported. More info `here`_.

        As a result, if a template isn't provided, the default template is
        modified to use a font named ``Courier New`` on Windows and
        ``Liberation Mono`` on other operating systems in the ``Normal``
        style.

    .. _here:
        https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html

    """
    blanks_needed = check_fit(plaintext, ciphertext)
    if blanks_needed > 0:
        raise ValueError(f'Need {blanks_needed} more spaces in the plaintext '
                         f'(fake) message.')
    if template_path is None:
        # Modify default template.
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        if system().lower().startswith('windows'):
            font.name = 'Courier New'
        else:
            font.name = 'Liberation Mono'
    else:
        doc = docx.Document(template_path)

    line_index, letter_index = 0, 0
    for line in plaintext:
        # Add new paragraph to template.
        paragraph = doc.add_paragraph()
        paragraph_index = len(doc.paragraphs) - 1
        for letter in line:
            # Add each letter to paragraph.
            if all([letter == ' ',
                    letter_index < len(ciphertext[line_index])]):
                # Add real message to space and set color to white.
                paragraph.add_run(ciphertext[line_index][letter_index])
                run = doc.paragraphs[paragraph_index].runs[-1]
                font = run.font
                # Make red for testing.
                font.color.rgb = RGBColor(255, 255, 255)
                letter_index += 1
            else:
                paragraph.add_run(letter)
            if all([letter_index >= len(ciphertext[line_index]),
                    line_index < len(ciphertext) - 1]):
                # Go to next line in ciphertext if end reached.
                line_index += 1
                letter_index = 0
    doc.save(filename)


def main(fakefile: str = None, cipherfile: str = None,
         savepath: str = None) -> None:
    """Demonstrate the invisible ink writer.

    Demonstrate :func:`write_invisible`, but for testing,
    it is a basic wrapper function for :func:`write_invisible`.
    Embed **cipherfile** in **fakefile**'s whitespace.

    Args:
        fakefile (str): Path to .docx file with fake message.
            Defaults to ``./c1files/fake.docx``.
        cipherfile (str): Path to .docx file with real message.
            Defaults to ``./c1files/real.docx``.
        savepath (str): Path to .docx file for output.
            Defaults to ``./c1files/Hello.docx``.

    Returns:
         :py:obj:`None`. The contents of **cipherfile**'s text is embedded
         in **fakefile**'s whitespace and saved to **savepath**.

    """
    print('I can embed a hidden message in a .docx file\'s white space '
          'between letters by making the font\ncolor white. It\'s far less '
          'bulletproof than it sounds.\n')
    current_dir = Path('./c1files').resolve()
    if fakefile is None or cipherfile is None:
        fakefile = PurePath(current_dir).joinpath('fake.docx')
        cipherfile = PurePath(current_dir).joinpath('real.docx')
    if savepath is None:
        savepath = PurePath(current_dir).joinpath('DearInternet.docx')
    faketext = get_text(fakefile, False)
    ciphertext = get_text(cipherfile)
    write_invisible(faketext, ciphertext, None, savepath)
    print('Fin.\n')
    print('To read the hidden message, select the entire document and\n'
          'highlight it a darkish gray.')


if __name__ == '__main__':
    main()
